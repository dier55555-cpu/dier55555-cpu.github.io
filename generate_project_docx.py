#!/usr/bin/env python3
"""Generate a DOCX project document: implementation of the TZ with Noya."""

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn, nsmap
from docx.shared import Cm, Pt, RGBColor, Emu


NAVY = RGBColor(0x1B, 0x3A, 0x5F)
ACCENT = RGBColor(0xC4, 0x5C, 0x26)
DARK = RGBColor(0x1A, 0x1A, 0x1A)
GRAY = RGBColor(0x4A, 0x4A, 0x4A)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
HEADER_BG = "1B3A5F"
ALT_ROW = "F4F1EC"
LIGHT_BG = "EEF3F8"


def set_run_font(run, name="Calibri", size=11, bold=False, italic=False, color=DARK):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color


def add_text(paragraph, text, **kwargs):
    run = paragraph.add_run(text)
    set_run_font(run, **kwargs)
    return run


def shade_cell(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def set_cell_borders(cell, color="C8C4BC", sz="4"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), sz)
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        tcBorders.append(el)
    tcPr.append(tcBorders)


def set_cell_text(cell, text, *, bold=False, color=DARK, size=10, align="left", fill=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
    }[align]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.15
    add_text(p, text, size=size, bold=bold, color=color)
    if fill:
        shade_cell(cell, fill)
    set_cell_borders(cell)
    for p_el in cell.paragraphs:
        p_el.paragraph_format.left_indent = Cm(0.1)
        p_el.paragraph_format.right_indent = Cm(0.1)


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "C8C4BC")
        borders.append(el)
    tblPr.append(borders)

    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True, color=WHITE, size=10, fill=HEADER_BG)

    for r_idx, row in enumerate(rows):
        fill = ALT_ROW if r_idx % 2 else "FFFFFF"
        for c_idx, val in enumerate(row):
            set_cell_text(table.rows[r_idx + 1].cells[c_idx], val, size=10, fill=fill)

    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return table


def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = NAVY if level == 1 else ACCENT if level == 2 else NAVY
        run.font.name = "Calibri"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    return p


def body(doc, text, *, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.2
    add_text(p, text, size=size, bold=bold, italic=italic)
    return p


def bullet(doc, text, *, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.clear()
    p.paragraph_format.left_indent = Cm(1.25 + level * 0.6)
    p.paragraph_format.space_after = Pt(3)
    add_text(p, text, size=11)
    return p


def numbered(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.clear()
    p.paragraph_format.space_after = Pt(3)
    add_text(p, text, size=11)
    return p


def callout(doc, title, text):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    cell.text = ""
    shade_cell(cell, LIGHT_BG)
    set_cell_borders(cell, color="1B3A5F", sz="8")
    p1 = cell.paragraphs[0]
    p1.paragraph_format.space_before = Pt(6)
    p1.paragraph_format.space_after = Pt(2)
    add_text(p1, title, size=11, bold=True, color=NAVY)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(6)
    add_text(p2, text, size=10, color=GRAY)
    doc.add_paragraph()


def code_block(doc, text):
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    cell.text = ""
    shade_cell(cell, "F7F5F2")
    set_cell_borders(cell, color="D4CFC6", sz="4")
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    run.font.size = Pt(9)
    run.font.color.rgb = DARK
    doc.add_paragraph()


def set_header_footer(doc):
    section = doc.sections[0]
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_text(hp, "Реализация ТЗ на платформе «Ноя»  ·  конфиденциально", size=9, italic=True, color=GRAY)

    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_text(fp, "Веб-приложение расчётов и коммерческих предложений  ·  стр. ", size=9, color=GRAY)

    # PAGE field
    run = fp.add_run()
    set_run_font(run, size=9, color=GRAY)
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr)
    run._r.append(fld_char_end)


def build():
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.0)
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)

    set_header_footer(doc)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(11)
    styles["Normal"].font.color.rgb = DARK
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")

    # ----- Title page -----
    for _ in range(3):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_text(p, "ПРОЕКТ РЕАЛИЗАЦИИ", size=14, bold=True, color=ACCENT)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    add_text(
        p,
        "Веб-приложение расчётов\nи коммерческих предложений\nс интеграцией amoCRM",
        size=26,
        bold=True,
        color=NAVY,
    )

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(18)
    add_text(p, "На платформе «Ноя» (noya-ai.ru) и её возможностях", size=14, italic=True, color=GRAY)

    doc.add_paragraph()

    meta = [
        ("Основание", "Приложение № 1 — Техническое задание (ТЗ)"),
        ("Платформа реализации", "Ноя-билдер: веб-приложение, база, роли, серверные функции, ИИ"),
        ("Интеграция", "amoCRM (карточка сделки: виджет / iframe + API)"),
        ("Калькуляторы", "№ 1 — реечный потолок;  № 2 — подбор светильников"),
        ("Дата", "22 августа 2026 г."),
        ("Статус", "Рабочий проект внедрения"),
    ]
    table = doc.add_table(rows=len(meta), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (k, v) in enumerate(meta):
        fill = ALT_ROW if i % 2 else "FFFFFF"
        set_cell_text(table.rows[i].cells[0], k, bold=True, color=NAVY, size=10, fill=fill)
        set_cell_text(table.rows[i].cells[1], v, size=10, fill=fill)
        table.rows[i].cells[0].width = Cm(5.5)
        table.rows[i].cells[1].width = Cm(11.0)

    doc.add_page_break()

    # ----- 1. Summary -----
    heading(doc, "1. Краткий вывод", 1)
    body(
        doc,
        "Техническое задание реализуется на платформе «Ноя» как полноценное веб-приложение "
        "с личным кабинетом, а не как чат-бот в мессенджерах. ИИ-агент Нои — вспомогательный "
        "контур: распознавание размеров с плана и синхронизация сделки. Ядро продукта — "
        "детерминированные формулы, справочники администратора и выгрузка коммерческого предложения.",
    )
    callout(
        doc,
        "Принцип разделения контуров",
        "Расчёт всегда выполняется кодом серверной функции по формулам ТЗ. "
        "ИИ используется только для извлечения длины и ширины с чертежа и для подсказок. "
        "Менеджер подтверждает или правит цифры вручную — как требует п. 2.1 ТЗ.",
    )
    body(
        doc,
        "Ноя закрывает порядка 80 % требований ТЗ штатными средствами билдера: роли, база, "
        "файлы, HTTPS, российская инфраструктура, vision/OCR, запись в amoCRM. Снаружи остаются "
        "тонкий виджет карточки сделки и, при жёстких требованиях к чертежам, специализированный OCR.",
    )

    heading(doc, "2. Соответствие ТЗ возможностям Нои", 1)
    add_table(
        doc,
        ["Требование ТЗ", "Чем закрывается в Ное"],
        [
            [
                "Веб-приложение, Chrome / Edge / Firefox",
                "React/TypeScript-приложение в Ноя-билдере; публикация на HTTPS и своём домене",
            ],
            [
                "ЛК, логин/пароль, роли Админ и Менеджер",
                "Авторизация и роли платформы; закрытые разделы админки и калькуляторов",
            ],
            [
                "Справочники, оргнастройки, пользователи",
                "База данных (Supabase на российских серверах) и CRUD-админка",
            ],
            [
                "Расчёты и формирование КП",
                "Формы, таблицы спецификации, серверные (edge) функции",
            ],
            [
                "Доступ из карточки сделки amoCRM",
                "Интеграция amoCRM (чтение/запись полей, воронка) + тонкий виджет с iframe",
            ],
            [
                "HTTPS",
                "Штатный SSL-сертификат при привязке домена",
            ],
            [
                "Загрузка плана JPG / PNG / PDF",
                "Файловое хранилище проекта и серверные функции",
            ],
            [
                "Распознавание размеров (CV / OCR)",
                "Встроенное распознавание фото и OCR + vision-модель внутри приложения",
            ],
            [
                "Экспорт в Excel и Word",
                "Клиентские библиотеки в React либо генерация файла на серверной функции",
            ],
            [
                "ИИ без зарубежных карт и VPN",
                "Ключи моделей платформы; инфраструктура в РФ; маскирование ПДн",
            ],
        ],
        col_widths=[7.5, 9.0],
    )

    heading(doc, "3. Целевая архитектура", 1)
    body(doc, "Система состоит из трёх слоёв.")
    numbered(
        doc,
        "Веб-приложение Нои — основной продукт: вход, роли, справочники, два калькулятора, история КП, экспорт.",
    )
    numbered(
        doc,
        "Тонкий виджет amoCRM — вкладка или правая колонка карточки сделки: iframe приложения с lead_id и подписью.",
    )
    numbered(
        doc,
        "Синхронизация — серверная функция пишет в сделку сумму, состав, файл КП и при необходимости двигает этап воронки.",
    )

    heading(doc, "3.1. Схема взаимодействия", 2)
    code_block(
        doc,
        "amoCRM  (карточка сделки)\n"
        "   │  виджет: вкладка «Расчёт» / правая колонка\n"
        "   │  iframe + lead_id + account_id + подпись\n"
        "   ▼\n"
        "Веб-приложение Ноя\n"
        "   /login\n"
        "   /admin              справочники, пользователи, нормы\n"
        "   /calc/ceiling       калькулятор № 1 — реечный потолок\n"
        "   /calc/lighting      калькулятор № 2 — светильники\n"
        "   /offers             история коммерческих предложений\n"
        "   │\n"
        "   ├── База Нои          справочники, расчёты, пользователи\n"
        "   ├── Файлы             планы помещений, файлы КП\n"
        "   └── ИИ Нои            разбор чертежа → JSON размеров\n"
        "          │\n"
        "          └── amoCRM API: сумма, спецификация, файл, этап",
    )

    heading(doc, "3.2. Маршруты и роли", 2)
    add_table(
        doc,
        ["Роль", "Доступ", "Запрещено"],
        [
            [
                "Администратор",
                "Справочники, нормы, пользователи, реквизиты компании, коэффициент по умолчанию, все расчёты",
                "—",
            ],
            [
                "Менеджер",
                "Калькуляторы, свои КП, экспорт Excel/Word, работа из карточки сделки",
                "Редактирование справочников, пользователей и норм",
            ],
        ],
        col_widths=[4.0, 7.0, 5.5],
    )
    body(
        doc,
        "Вход в отдельном окне — email и пароль Нои. Из amoCRM предпочтителен SSO: виджет передаёт "
        "идентификатор сделки и подпись, серверная функция проверяет секрет интеграции и открывает "
        "сессию менеджера без повторного логина.",
    )

    heading(doc, "4. Реализация по разделам ТЗ", 1)

    heading(doc, "4.1. Общие положения (п. 1 ТЗ)", 2)
    add_table(
        doc,
        ["Пункт", "Решение"],
        [
            [
                "1.1  Платформа",
                "Одно веб-приложение в Ноя-билдере. Публикация на своём домене, адаптив под десктоп менеджера.",
            ],
            [
                "1.2  Аутентификация и роли",
                "Логин/пароль. Две роли: Администратор и Менеджер. Админ ведёт справочники и пользователей.",
            ],
            [
                "1.3  CRM",
                "Приложение открывается из карточки сделки amoCRM (iframe-виджет). Расчёт привязан к lead_id.",
            ],
            [
                "1.4  Безопасность",
                "HTTPS, секреты в шифрованном хранилище Нои, маскирование ПДн при вызовах модели.",
            ],
        ],
        col_widths=[4.5, 12.0],
    )
    body(
        doc,
        "Готовой кнопки «встроить приложение Нои во вкладку сделки» нет. Виджет amoCRM — отдельный "
        "тонкий архив (JS + manifest.json). Само приложение в SDK amoCRM не переносится.",
    )
    body(
        doc,
        "Необходимо явно разрешить встраивание: CSP frame-ancestors для доменов *.amocrm.ru "
        "(и при необходимости kommo.com). Без этого iframe в карточке сделки будет заблокирован браузером.",
    )

    heading(doc, "4.2. Калькулятор № 1 — реечный потолок (п. 2 ТЗ)", 2)

    heading(doc, "Ввод данных менеджером", 3)
    bullet(doc, "Загрузка плана помещения: JPG, PNG, PDF. PDF — рендер первой страницы в изображение.")
    bullet(doc, "Кнопка «Распознать размеры»: серверная функция вызывает vision/OCR Нои.")
    bullet(doc, "Модель возвращает только JSON: length_m, width_m, confidence, notes.")
    bullet(doc, "При низкой уверенности поля остаются пустыми — менеджер вводит длину и ширину вручную.")
    bullet(doc, "Размер рейки, шаг, материал — из справочников администратора; шаг также можно ввести вручную.")
    bullet(doc, "Цена материала — из справочника (руб./м² или руб./шт.) либо ручной ввод.")
    bullet(doc, "Коэффициент формирования стоимости: слайдер и число, диапазон 1,5–2,5, шаг 0,1.")

    heading(doc, "Промпт распознавания плана (контур ИИ)", 3)
    body(
        doc,
        "ИИ не считает спецификацию. Его единственная задача — вытащить подписанные размеры с чертежа.",
    )
    code_block(
        doc,
        "Роль: извлечение размеров помещения с плана (чертёж / скан / PDF).\n"
        "Верни ТОЛЬКО JSON без markdown:\n"
        "{\n"
        '  "length_m": number | null,\n'
        '  "width_m": number | null,\n'
        '  "unit_detected": "m" | "mm" | "cm" | null,\n'
        '  "confidence": 0.0-1.0,\n'
        '  "notes": "что не удалось прочитать"\n'
        "}\n"
        "Правила:\n"
        "— бери только явно подписанные габариты помещения;\n"
        "— не вычисляй площадь по пикселям и масштабу, если размеры не подписаны;\n"
        "— если сомневаешься — null и низкий confidence;\n"
        "— не предлагай спецификацию, цены и материалы.",
    )

    heading(doc, "Алгоритм расчёта (только код)", 3)
    body(doc, "Все нормы и справочные величины редактирует администратор. Формулы соответствуют п. 2.2 ТЗ:")
    code_block(
        doc,
        "S            = длина × ширина\n"
        "периметр     = 2 × (длина + ширина)\n"
        "\n"
        "рейки        = ceil( (S / (ширина_рейки + шаг)) × 1.05 )\n"
        "заглушки     = периметр / длина_одной_заглушки\n"
        "стрингеры    = (длина / шаг_стрингера) × число_рядов_по_ширине\n"
        "подвесы      = S × норма_шт_на_м²\n"
        "крепеж       = f(стрингеры, подвесы)     // нормы из справочника\n"
        "\n"
        "материалы    = Σ (количество × цена)\n"
        "комплектующ. = Σ (количество × цена)\n"
        "итог         = (материалы + комплектующие) × коэффициент",
    )
    body(
        doc,
        "Формулы целесообразно хранить в таблице «Правила расчёта», чтобы менять коэффициенты запаса "
        "и нормы без правки кода приложения.",
    )

    heading(doc, "Выходные данные", 3)
    bullet(doc, "Таблица спецификации: наименование, единица, количество, цена за единицу, сумма.")
    bullet(doc, "Итоговая стоимость по формуле ТЗ с выбранным коэффициентом.")
    bullet(doc, "Экспорт в Excel (.xlsx) и Word (.docx) по шаблону компании: логотип, реквизиты, номер КП, сделка.")
    bullet(doc, "Сохранение расчёта с привязкой к пользователю и lead_id amoCRM.")

    heading(doc, "4.3. Калькулятор № 2 — светильники (п. 3 ТЗ)", 2)
    body(
        doc,
        "В исходном ТЗ для калькулятора № 2 приведены входные данные; алгоритм подбора и формат "
        "выхода в приложении обрезаны. Ниже — рабочая схема, которую нужно подтвердить с заказчиком "
        "до разработки.",
    )
    add_table(
        doc,
        ["Ввод менеджера (по ТЗ)", "Как реализуется"],
        [
            ["Габариты L × W × H, мм", "Числовые поля с валидацией"],
            ["Мощность, Вт; световой поток, лм", "Числовые поля"],
            ["Цвет / тип диода", "Список: 3000K, 4000K, 5000K, RGB, RGBW"],
            ["Угол излучения, °", "Список 15 / 30 / 45 / 60 / 90 / 120 или ручной ввод"],
        ],
        col_widths=[7.5, 9.0],
    )
    bullet(doc, "Справочник светильников: габариты, мощность, поток, цвет, угол, цена, наличие.")
    bullet(doc, "Фильтр + скоринг: допуск по мм, ближайшие поток и мощность, точное совпадение цвета и угла.")
    bullet(doc, "Если точного совпадения нет — сообщение менеджеру и ручной выбор из отфильтрованного списка.")
    bullet(doc, "ИИ не подбирает изделие сам: может только пояснить, почему позиция отсеялась.")
    bullet(doc, "Дальше — та же спецификация, коэффициент (если применим) и выгрузка КП.")
    callout(
        doc,
        "Нужно уточнить у заказчика",
        "Допуски по габаритам; приоритет полей (поток важнее мощности или наоборот); "
        "можно ли предлагать ближайший угол/CCT; что делать при отсутствии позиции; "
        "нужен ли коэффициент наценки как в калькуляторе № 1; состав спецификации (крепёж, БП, кабель).",
    )

    heading(doc, "5. Модель данных (справочники и расчёты)", 1)
    body(doc, "База поднимается агентом Нои. Ниже — минимальный набор сущностей.")

    heading(doc, "5.1. Справочники администратора", 2)
    add_table(
        doc,
        ["Сущность", "Ключевые поля"],
        [
            ["users", "id, email, ФИО, роль (admin/manager), активен, amo_user_id"],
            ["organizations", "реквизиты, логотип, шаблон КП, коэффициент по умолчанию"],
            ["slat_sizes", "наименование, ширина мм, длина м, материал_id, цена, ед. изм."],
            ["materials", "алюминий / сталь / ПВХ и др., плотность/описание"],
            ["steps", "шаг между рейками, мм; признак «из списка / произвольный»"],
            ["end_caps", "наименование, длина шт., цена"],
            ["stringer_norms", "шаг стрингера, правило рядов по ширине"],
            ["hanger_norms", "норма шт/м²"],
            ["fastener_norms", "норма на стрингер / на подвес, цена"],
            ["calc_rules", "коэффициент запаса реек (1,05), округление, формулы"],
            ["luminaires", "L, W, H, Вт, лм, CCT/тип, угол, цена, наличие, артикул"],
            ["price_lists", "версия прайса, дата, источник"],
        ],
        col_widths=[4.5, 12.0],
    )

    heading(doc, "5.2. Операционные данные", 2)
    add_table(
        doc,
        ["Сущность", "Ключевые поля"],
        [
            [
                "calculations",
                "id, type (ceiling/lighting), manager_id, lead_id, входы, коэффициент, итог, статус",
            ],
            [
                "calculation_items",
                "calculation_id, наименование, ед., количество, цена, сумма, группа (материал/комплектующие)",
            ],
            ["plan_files", "calculation_id, файл, результат распознавания JSON, кто подтвердил"],
            ["offers", "calculation_id, номер КП, xlsx, docx, дата, отправлено_в_amo"],
        ],
        col_widths=[4.5, 12.0],
    )

    heading(doc, "6. Интеграция с amoCRM", 1)
    heading(doc, "6.1. Два контура", 2)
    numbered(
        doc,
        "Встраивание UI. Виджет amoCRM (locations: lcard, card_sdk или правая колонка lcard-1) "
        "рисует iframe на URL приложения: https://calc.domain.ru/embed?lead_id=&account_id=&sig=",
    )
    numbered(
        doc,
        "Данные сделки. Штатная интеграция Нои и/или серверная функция по API amoCRM: "
        "чтение названия и полей сделки, запись суммы, ссылки на КП, файла, смена этапа.",
    )

    heading(doc, "6.2. Поток работы менеджера", 2)
    numbered(doc, "Открывает карточку сделки в amoCRM.")
    numbered(doc, "Переходит во вкладку «Расчёт» — загружается приложение Нои с контекстом сделки.")
    numbered(doc, "Выбирает калькулятор, загружает план или вводит размеры, заполняет параметры.")
    numbered(doc, "Получает спецификацию, двигает коэффициент, выгружает Excel/Word.")
    numbered(doc, "Нажимает «Сохранить в сделку» — итог и файл уходят в amoCRM.")

    heading(doc, "6.3. Что пишем в сделку", 2)
    add_table(
        doc,
        ["Поле / сущность amoCRM", "Содержимое"],
        [
            ["Сумма сделки", "Итог расчёта (материалы + комплектующие) × коэффициент"],
            ["Кастомное поле «Коэффициент»", "Значение 1,5–2,5"],
            ["Кастомное поле «Тип расчёта»", "Реечный потолок / Светильники"],
            ["Примечание или файл", "PDF/DOCX/XLSX коммерческого предложения"],
            ["Ссылка", "URL расчёта в приложении Нои"],
            ["Этап воронки (опционально)", "Например, «КП отправлено»"],
        ],
        col_widths=[7.0, 9.5],
    )

    heading(doc, "7. Экспорт Excel и Word", 1)
    body(
        doc,
        "Генерация документов выполняется в самом приложении (клиентские библиотеки в React "
        "или серверная функция). Шаблоны хранит администратор.",
    )
    add_table(
        doc,
        ["Формат", "Состав файла"],
        [
            [
                "Excel (.xlsx)",
                "Лист «Спецификация»: колонки ТЗ; лист «Итог»: коэффициент, суммы, реквизиты, lead_id",
            ],
            [
                "Word (.docx)",
                "Титул КП, реквизиты сторон, таблица спецификации, итог прописью, комментарий менеджера",
            ],
        ],
        col_widths=[4.5, 12.0],
    )
    body(
        doc,
        "Номер КП формируется в приложении (префикс организации + год + порядковый номер) "
        "и дублируется в имя файла и в поле сделки.",
    )

    heading(doc, "8. Ограничения Нои и внешние работы", 1)
    add_table(
        doc,
        ["Ограничение", "Как закрываем"],
        [
            [
                "Нет нативного виджета «вкладка сделки»",
                "Отдельный тонкий виджет amoCRM с iframe; при необходимости — кастомная интеграция через команду Нои",
            ],
            [
                "Vision/OCR не равен промышленному CV чертежей",
                "Распознавание — ускоритель ввода. Истина — ручное подтверждение. При жёстком SLA — внешний OCR (например Smart Engines) по HTTP",
            ],
            [
                "ИИ может «посчитать» спецификацию неверно",
                "Запрет: формулы только в коде и справочниках. Тест-кейсы на типовые комнаты",
            ],
            [
                "Платформа не рассчитана на самохостинг",
                "Приложение живёт в Ное; домен клиента, данные в РФ",
            ],
            [
                "Нет нативных iOS/Android",
                "Не требуется ТЗ: веб в браузере менеджера",
            ],
            [
                "Неполный п. 3 ТЗ",
                "Зафиксировать правила подбора светильников до этапа 5",
            ],
        ],
        col_widths=[6.5, 10.0],
    )

    heading(doc, "9. План внедрения", 1)
    body(
        doc,
        "Этапы независимы по риску распознавания и неполного ТЗ по светильникам. "
        "После этапа 2 менеджер уже считает потолок и выгружает КП без amoCRM и без CV.",
    )
    add_table(
        doc,
        ["Этап", "Состав", "Результат"],
        [
            [
                "1. Каркас ЛК",
                "Авторизация, роли, справочники реек, материалов, норм, крепежа, светильников, карточка организации",
                "Админ заполняет прайс и нормы",
            ],
            [
                "2. Калькулятор № 1 без CV",
                "Ручной ввод L×W, формулы, спецификация, сохранение, Excel и Word",
                "Рабочий контур менеджера",
            ],
            [
                "3. amoCRM",
                "Виджет-iframe, прокидывание lead_id, SSO, запись суммы и файла в сделку",
                "Расчёт из карточки сделки",
            ],
            [
                "4. Распознавание плана",
                "Загрузка JPG/PNG/PDF, vision, автозаполнение размеров, обязательное подтверждение",
                "Ускорение ввода, ручной fallback",
            ],
            [
                "5. Калькулятор № 2",
                "После фиксации правил подбора: фильтр, скоринг, спецификация, КП",
                "Второй калькулятор в том же ЛК",
            ],
            [
                "6. ИИ-агент (опционально)",
                "Дожим КП, запрос «пересчитай сделку N с коэффициентом 2.0» из Telegram",
                "Сервис поверх готового расчётного ядра",
            ],
        ],
        col_widths=[3.5, 7.5, 5.5],
    )

    heading(doc, "10. Как ставить задачу агенту-билдеру Нои", 1)
    body(
        doc,
        "Не скармливать всё ТЗ одной фразой. Идти итерациями. Формулы и поля справочников "
        "прикладывать таблицами. После генерации расчёта — прогнать контрольные примеры.",
    )
    numbered(doc, "«Собери веб-приложение: ЛК, роли admin/manager, справочники с полями из раздела 5».")
    numbered(
        doc,
        "«Калькулятор реечного потолка, формулы строго такие… коэффициент 1,5–2,5 шаг 0,1».",
    )
    numbered(doc, "«Экспорт спецификации в xlsx и docx по шаблону компании».")
    numbered(doc, "«Загрузка плана, вызов vision, JSON размеров, ручное подтверждение».")
    numbered(doc, "«Эндпоинт: по lead_id отдать и принять расчёт, записать итог в amoCRM».")

    heading(doc, "10.1. Контрольный пример (потолок)", 2)
    body(doc, "Использовать для приёмки формул, не для дизайна.")
    add_table(
        doc,
        ["Параметр", "Значение"],
        [
            ["Помещение", "6,00 × 4,00 м  (S = 24 м², периметр = 20 м)"],
            ["Ширина рейки + шаг", "из справочника, например 100 мм + 20 мм = 0,12 м"],
            ["Коэффициент запаса реек", "1,05"],
            ["Коэффициент цены", "2,0"],
            ["Ожидание", "количества = формулы п. 2.2; итог = (материалы + комплектующие) × 2,0"],
        ],
        col_widths=[5.5, 11.0],
    )

    heading(doc, "11. Состав работ вне билдера", 1)
    bullet(doc, "Регистрация приватной интеграции amoCRM, виджет, manifest, области lcard / card_sdk.")
    bullet(doc, "Кастомные поля сделки: коэффициент, тип расчёта, ссылка на КП.")
    bullet(doc, "Секрет подписи iframe и проверка на серверной функции.")
    bullet(doc, "Шаблоны Word/Excel с фирменным стилем заказчика.")
    bullet(doc, "Наполнение справочников реальными ценами и нормами.")
    bullet(doc, "При необходимости — договорённость с командой Нои на frame-ancestors и кастомную интеграцию.")

    heading(doc, "12. Итог", 1)
    body(
        doc,
        "Проект соответствует ТЗ и реализуется на «Ное» как веб-приложение с личным кабинетом. "
        "Платформа закрывает роли, справочники, расчёты, файлы, HTTPS, ИИ для чертежа и выгрузку "
        "в amoCRM. Снаружи остаются тонкий виджет карточки сделки и опциональный промышленный OCR.",
    )
    body(
        doc,
        "Расчёт — в коде. ИИ — только читает план и помогает в диалоге. "
        "Первый полезный результат для менеджера появляется после этапа 2; полный контур ТЗ — "
        "после этапов 3–5.",
    )

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    add_text(p, "Документ подготовлен как рабочий проект внедрения на основании Приложения № 1 (ТЗ).", italic=True, size=10, color=GRAY)

    out = "/workspace/docs/Proekt_realizacii_TZ_Noya.docx"
    import os

    os.makedirs("/workspace/docs", exist_ok=True)
    doc.save(out)
    print(out)


if __name__ == "__main__":
    build()
